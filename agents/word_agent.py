"""
Word Agent — Creates and edits Word documents (.docx).
Uses python-docx for document creation and win32com for PDF export.
"""

import os
import logging
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from agents.base_agent import BaseAgent, AgentResult
from config import OUTPUT_DIR

logger = logging.getLogger("word_agent")


class WordAgent(BaseAgent):
    name = "word"
    description = "Create and edit Word documents (.docx), add tables, export to PDF"
    capabilities = ["create_document", "add_table", "export_pdf"]

    def execute(self, action: str, args: dict) -> AgentResult:
        self._validate_action(action)

        if action == "create_document":
            return self._create_document(**args)
        elif action == "add_table":
            return self._add_table(**args)
        elif action == "export_pdf":
            return self._export_pdf(**args)

        return AgentResult(success=False, message=f"Unknown action: {action}")

    # ── Actions ──────────────────────────────

    def _create_document(self, filename: str, title: str, content: str) -> AgentResult:
        """Create a formatted .docx Word document."""
        try:
            if not filename:
                import time
                filename = f"document_{int(time.time())}.docx"
            if not title:
                title = "Generated Document"
            if not content:
                content = ""

            doc = Document()

            # Title
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

            # Timestamp
            ts_para = doc.add_paragraph(
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            )
            ts_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in ts_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            doc.add_paragraph("")  # spacer

            # Parse content — supports markdown-like headings
            for line in content.strip().split("\n"):
                if line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.strip().startswith("- ") or line.strip().startswith("• "):
                    # Bullet point
                    doc.add_paragraph(line.strip()[2:], style="List Bullet")
                elif line.strip() == "":
                    doc.add_paragraph("")
                else:
                    doc.add_paragraph(line)

            # Save
            if not filename.endswith(".docx"):
                filename += ".docx"
            path = os.path.join(OUTPUT_DIR, filename)
            doc.save(path)

            logger.info("Created Word document: %s", path)
            return AgentResult(
                success=True,
                message=f"✅ Word document created: {filename}",
                output_file=path
            )

        except Exception as e:
            logger.exception("Failed to create Word document")
            return AgentResult(success=False, message=f"❌ Word error: {e}")

    def _add_table(self, filename: str, headers: list, rows: list) -> AgentResult:
        """Add a table to an existing or new Word document."""
        try:
            if not filename.endswith(".docx"):
                filename += ".docx"
            path = os.path.join(OUTPUT_DIR, filename)

            # Open existing or create new
            if os.path.exists(path):
                doc = Document(path)
            else:
                doc = Document()
                doc.add_heading(filename.replace(".docx", "").replace("_", " ").title(), level=0)

            # Create table
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Light Grid Accent 1"

            # Headers
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = str(header)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)

            # Data rows
            for row_idx, row_data in enumerate(rows):
                for col_idx, value in enumerate(row_data):
                    table.rows[row_idx + 1].cells[col_idx].text = str(value)

            doc.save(path)
            logger.info("Added table to: %s", path)
            return AgentResult(
                success=True,
                message=f"✅ Table added to {filename}",
                output_file=path
            )

        except Exception as e:
            logger.exception("Failed to add table")
            return AgentResult(success=False, message=f"❌ Table error: {e}")

    def _export_pdf(self, filename: str) -> AgentResult:
        """Export a .docx to PDF using Word COM automation."""
        try:
            if not filename.endswith(".docx"):
                filename += ".docx"
            docx_path = os.path.join(OUTPUT_DIR, filename)

            if not os.path.exists(docx_path):
                return AgentResult(success=False, message=f"❌ File not found: {filename}")

            pdf_path = docx_path.replace(".docx", ".pdf")

            # Use win32com to open Word and save as PDF
            import win32com.client
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            try:
                doc = word_app.Documents.Open(os.path.abspath(docx_path))
                doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = PDF
                doc.Close()
            finally:
                word_app.Quit()

            logger.info("Exported PDF: %s", pdf_path)
            return AgentResult(
                success=True,
                message=f"✅ Exported to PDF: {os.path.basename(pdf_path)}",
                output_file=pdf_path
            )

        except ImportError:
            return AgentResult(
                success=False,
                message="❌ PDF export requires pywin32. Install with: pip install pywin32"
            )
        except Exception as e:
            logger.exception("PDF export failed")
            return AgentResult(success=False, message=f"❌ PDF export error: {e}")
